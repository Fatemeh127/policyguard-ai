"""Production-ready DOCX text extractor using python-docx (mypy-safe)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def load_docx(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc: DocxDocument = Document(str(path))

        text_parts: list[str] = []
        seen: set[str] = set()

        def add_text(text: Any) -> None:
            # Defensive typing because python-docx returns Any
            if text is None:
                return

            cleaned = str(text).strip()
            if not cleaned:
                return

            if cleaned in seen:
                return

            seen.add(cleaned)
            text_parts.append(cleaned)

        # 1. Paragraphs
        for p in doc.paragraphs:
            paragraph: Paragraph = p
            add_text(paragraph.text)

        # 2. Tables
        for table in doc.tables:
            t: Table = table
            for row in t.rows:
                for cell in row.cells:
                    c: _Cell = cell
                    add_text(c.text)

        result = "\n".join(text_parts).strip()

        if not result:
            raise ValueError("No extractable text found in DOCX (may contain only images/shapes).")

        return result

    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {e}") from e
